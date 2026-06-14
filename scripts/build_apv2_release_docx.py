from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]

SOURCES = [
    (
        ROOT / "docs" / "Release_APV2_FinalPaper_20260614.md",
        ROOT / "paper_artifacts" / "release_20260614" / "APV2_论文_白箱预测行动闭环架构_20260614.docx",
    ),
    (
        ROOT / "docs" / "Release_APV2_NewsArticle_20260614.md",
        ROOT / "paper_artifacts" / "release_20260614" / "APV2_新闻稿_让机器一直是它自己_20260614.docx",
    ),
    (
        ROOT / "docs" / "Release_APV2_RepositorySetup_20260614.md",
        ROOT / "paper_artifacts" / "release_20260614" / "APV2_仓库发布说明_20260614.docx",
    ),
]


def _set_run_font(run, *, name: str = "Microsoft YaHei", size: Pt = Pt(10.5), code: bool = False) -> None:
    font_name = "Consolas" if code else name
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = size


def _add_hyperlink(paragraph, text: str, url: str, *, size: Pt = Pt(10.5)) -> None:
    rel_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_inline_text(paragraph, text: str, *, size: Pt = Pt(10.5)) -> None:
    # Render common Markdown inline code spans cleanly in Word instead of
    # leaking literal backticks into the public document.
    pattern = re.compile(r"`([^`]+)`")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            _set_run_font(run, size=size)
        code_text = match.group(1)
        if code_text.startswith(("http://", "https://")):
            _add_hyperlink(paragraph, code_text, code_text, size=size)
        else:
            run = paragraph.add_run(code_text)
            _set_run_font(run, size=size, code=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_font(run, size=size)


def _set_cell_text(cell, text: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _add_inline_text(paragraph, text, size=Pt(9))


def _style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(len(table.columns)):
            text = row[c_idx].strip() if c_idx < len(row) else ""
            _set_cell_text(table.cell(r_idx, c_idx), text)
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _try_add_image(doc: Document, md_path: Path, line: str) -> bool:
    match = re.match(r"!\[(.*?)\]\((.*?)\)", line.strip())
    if not match:
        return False
    alt, raw_path = match.groups()
    image_path = Path(raw_path)
    if not image_path.is_absolute():
        image_path = (md_path.parent / image_path).resolve()
    if not image_path.exists():
        doc.add_paragraph(f"[Image not found: {raw_path}]")
        return True
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.9))
    if alt:
        cap = doc.add_paragraph(alt)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def _flush_table(doc: Document, table_lines: list[str]) -> None:
    if not table_lines:
        return
    rows: list[list[str]] = []
    for line in table_lines:
        parts = [p.strip() for p in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", p or "") for p in parts):
            continue
        rows.append(parts)
    _add_table(doc, rows)


def build_docx(md_path: Path, out_path: Path) -> None:
    doc = Document()
    _style_document(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    table_lines: list[str] = []
    in_code = False
    code_buffer: list[str] = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                run.font.size = Pt(9)
                code_buffer = []
                in_code = False
            else:
                _flush_table(doc, table_lines)
                table_lines = []
                in_code = True
            continue
        if in_code:
            code_buffer.append(line)
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            continue
        if table_lines:
            _flush_table(doc, table_lines)
            table_lines = []

        if not line.strip():
            continue
        if _try_add_image(doc, md_path, line):
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            _add_inline_text(p, line[2:].strip(), size=Pt(18))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            _add_inline_text(p, line[3:].strip(), size=Pt(16))
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            _add_inline_text(p, line[4:].strip(), size=Pt(14))
        elif line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            _add_inline_text(p, line[5:].strip(), size=Pt(12))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_text(p, line[2:].strip())
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline_text(p, re.sub(r"^\d+\. ", "", line).strip())
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(21)
            _add_inline_text(p, line)

    if table_lines:
        _flush_table(doc, table_lines)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    for src, out in SOURCES:
        build_docx(src, out)
        print(out)


if __name__ == "__main__":
    main()
